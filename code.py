from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os, subprocess, textwrap, re, math, datetime, pathlib, json

# ---------- Content (Paper 1: AQA 7552 Product Design, Section 3.1) ----------

guide = [
    {
        "section": "3.1.1",
        "title": "Materials and their applications",
        "topics": [
            {
                "title": "Selecting materials and justifying choices",
                "definitions": [
                    ("Physical properties", "Things you can measure/observe without changing the material (eg density, conductivity)."),
                    ("Mechanical properties", "How the material behaves under forces (eg tensile strength, toughness, hardness)."),
                    ("Working characteristics", "How the material behaves when you process or use it (eg formability, corrosion resistance)."),
                    ("Malleability", "How easily a material can be pressed/hammered into shape without cracking."),
                    ("Toughness", "Ability to absorb impact energy without snapping (impact resistance)."),
                    ("Hardness", "Resistance to scratching/indentation/wear."),
                    ("Corrosion / degradation", "Breaking down due to the environment (rusting, UV damage, chemical attack)."),
                    ("Thermal conductivity", "How well heat passes through a material."),
                    ("Electrical conductivity", "How well electricity passes through a material."),
                    ("Aesthetics", "How it looks/feels: colour, texture, finish, perceived quality."),
                    ("Manufacture", "How it can realistically be made at the required scale (prototype, batch, mass)."),
                    ("Disposal / end-of-life", "What happens after use: repair, reuse, recycle, landfill; ease of separation.")
                ],
                "explanations": [
                    "In Paper 1, you do not get marks for just naming a material. You get marks for explaining why it fits the job.",
                    "A strong way to justify material choice is: (1) what the part must do (function), (2) what forces and environment it faces, (3) what process and scale you will manufacture at, (4) how it should look/feel, and (5) what happens at end-of-life.",
                    "Example logic: a phone case needs toughness (survive drops), a little flexibility (snap on/off), decent surface finish (aesthetics), and a process like injection moulding for mass production. That pushes you toward a tough thermoplastic rather than a brittle or heat-set material.",
                    "The specification also expects you to do simple material quantity/cost calculations. This often means calculating sheet area/volume, adding waste allowance, multiplying by cost per unit, and comparing options."
                ],
                "pitfalls": [
                    "Listing properties with no link to the product (eg “ABS is strong” but not saying why strength matters here).",
                    "Mixing up hardness and toughness (hard can still be brittle).",
                    "Ignoring manufacture (choosing a material that cannot be made by the stated process/scale).",
                    "Forgetting end-of-life (mixed materials, coatings and adhesives can stop recycling).",
                    "Only talking about ‘cheap/expensive’ without mentioning tooling, waste, labour, or production scale."
                ],
                "exam_application": {
                    "prompt": "A handheld product casing must be light, impact resistant and suitable for mass production. Recommend a material and justify your choice.",
                    "what_to_do": [
                        "Name a suitable material (eg ABS, HIPS, or glass-filled nylon depending on loads).",
                        "Link properties to needs: toughness/impact resistance, low density, surface finish, durability.",
                        "Link to manufacture: injection moulding, repeatability, textured moulds, low unit cost at scale.",
                        "Add one sustainability/end-of-life point: recyclability, reducing part count, snap-fits for disassembly."
                    ]
                }
            },
            {
                "title": "Classification of materials",
                "definitions": [
                    ("Ferrous metal", "Contains iron; often strong but can rust (needs protection)."),
                    ("Non-ferrous metal", "No iron; often lighter and more corrosion resistant (eg aluminium, copper)."),
                    ("Alloy", "A mix of metals (or metal + non-metal) to improve properties (eg stainless steel)."),
                    ("Hardwood", "From deciduous trees; often denser (eg oak, beech)."),
                    ("Softwood", "From coniferous trees; often faster growing (eg pine, spruce)."),
                    ("Manufactured board", "Engineered wood sheet material (eg MDF, plywood, chipboard)."),
                    ("Thermoplastic", "Softens when heated and can be reshaped; often recyclable by remelting."),
                    ("Thermoset polymer", "Sets permanently when formed; does not remelt (usually better heat resistance)."),
                    ("Elastomer", "Rubber-like polymer that stretches and returns to shape (eg silicone, TPE)."),
                    ("Composite", "Two or more materials combined (matrix + reinforcement) to improve performance."),
                    ("Smart material", "Material that changes in response to a stimulus (heat, light, pressure)."),
                    ("Modern material", "New/advanced material developed for improved performance or new uses.")
                ],
                "explanations": [
                    "Classification helps you predict behaviour and suitable processes. For example, calling something a thermoplastic hints it can be moulded and potentially recycled; calling something ferrous hints you need to think about corrosion protection.",
                    "In exam answers, classification is often the ‘bridge’ between the material choice and the reasoning: you classify it, then explain what that classification implies for properties, processes, cost and end-of-life.",
                    "A common high-mark move is to mention combinations: a product might use a metal for strength, an elastomer for grip, and a coating for appearance/corrosion resistance."
                ],
                "pitfalls": [
                    "Calling stainless steel non-ferrous (it is an iron-based alloy, so ferrous).",
                    "Calling MDF a hardwood (it is a manufactured board).",
                    "Mixing up thermoplastic and thermoset (remeltable vs permanently set).",
                    "Giving ‘examples’ that are actually processes (eg ‘injection moulding’) instead of materials."
                ],
                "exam_application": {
                    "prompt": "Classify the material used for a drinks bottle and explain one advantage of that classification.",
                    "what_to_do": [
                        "Classify: ‘thermoplastic polymer’.",
                        "Advantage linked to classification: can be blow moulded for mass production; can be recycled by remelting (if collected and sorted)."
                    ]
                }
            },
            {
                "title": "Investigating and testing materials (and using data)",
                "definitions": [
                    ("Tensile strength", "Maximum pulling (tensile) stress a material can take before failure."),
                    ("Yield point", "Point where a material starts to permanently deform."),
                    ("Impact test / toughness", "Test that measures energy absorbed before fracture."),
                    ("Hardness test", "Test measuring resistance to indentation/scratching."),
                    ("Corrosion test", "Exposure test to see how fast a material degrades in a given environment."),
                    ("Conductivity test", "Measuring electrical resistance or heat transfer rate.")
                ],
                "explanations": [
                    "You should be able to describe what the test does, what is measured, and what the results mean for a product.",
                    "Tensile tests clamp a standard sample and pull it until it yields/fractures. Impact tests strike a sample and measure energy absorbed. Hardness tests press an indenter into the surface and measure indentation. Corrosion tests expose samples to moisture/salt/chemicals for a set time.",
                    "The key exam skill is analysing the results: choose the material that best matches the product need, and mention trade-offs. For example, a very hard material might scratch less but could be brittle and fail on impact."
                ],
                "pitfalls": [
                    "Describing the test but not saying what is measured (force, energy, indentation size, resistance).",
                    "Picking the ‘highest number’ without linking to the actual requirement.",
                    "Confusing strength and toughness (strength is resisting force; toughness is absorbing energy).",
                    "Not mentioning trade-offs when two materials each win in different properties."
                ],
                "exam_application": {
                    "prompt": "Two polymers are tested for impact energy absorbed and tensile strength. Recommend one for a phone case using the data.",
                    "what_to_do": [
                        "Prioritise toughness/impact for a phone case (drops).",
                        "Use numbers to compare (eg ‘A absorbs 2x the impact energy of B’).",
                        "Mention any trade-off (eg slightly lower tensile strength but much higher impact resistance).",
                        "Add a manufacturing note if relevant (eg injection moulding, surface finish)."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.2",
        "title": "Performance characteristics of materials",
        "topics": [
            {
                "title": "Papers and boards: characteristics + uses",
                "definitions": [
                    ("Score", "A controlled ‘weak line’ so card can fold neatly without cracking."),
                    ("Grain direction", "Direction fibres mostly run; affects strength and folding (folding against grain can crack)."),
                    ("Surface finish", "Texture/coating that affects printing quality and appearance."),
                    ("Biodegradable", "Can break down naturally by microorganisms (under the right conditions)."),
                    ("Recyclable", "Can be processed into new material (clean, sorted and uncontaminated).")
                ],
                "explanations": [
                    "Paper and board are used heavily in packaging because they are light, easy to cut/score/fold, and often recyclable.",
                    "Performance characteristics you need to talk about include: ability to be scored, cut, folded, and the surface qualities for printing. In real products, grain direction and thickness also control fold quality and stiffness.",
                    "Common paper/board types and typical uses (as per the spec): layout paper (sketch pads), cartridge paper (printing), tracing paper (copying images), bleed-proof paper (marker rendering), treated paper (photographic printing), watercolour paper (painting), corrugated card (packaging), bleached card (greeting cards/high quality packaging), mount board (modelling), duplex card (food packaging), foil-backed/laminated card (drinks packaging), metal effect card (gift packaging), moulded paper pulp (eco-friendly packaging)."
                ],
                "pitfalls": [
                    "Forgetting that folding thick card usually needs scoring (otherwise it cracks).",
                    "Ignoring recyclability rules: coatings/laminates can make recycling harder.",
                    "Not linking the paper type to the job (eg using tracing paper for a load-bearing package)."
                ],
                "exam_application": {
                    "prompt": "A company wants eco-friendly drinks packaging. Recommend a paper/board material and one finishing choice, with reasons.",
                    "what_to_do": [
                        "Pick a suitable board (eg corrugated card or moulded paper pulp).",
                        "Explain: stiffness/protection, printability, recyclability/biodegradability.",
                        "Finishing: choose minimal coatings; mention how laminates can affect recycling."
                    ]
                }
            },
            {
                "title": "Polymer-based sheet and film: characteristics + uses",
                "definitions": [
                    ("Transparency", "You can see through clearly."),
                    ("Translucency", "Light passes through but you cannot see clearly."),
                    ("Thermoforming", "Heating a plastic sheet then forming it over/into a mould."),
                    ("Flexibility", "Bends without cracking and springs back (within limits).")
                ],
                "explanations": [
                    "Polymer sheet/film is used when you need light weight, moisture resistance, and fast processing (often for packaging and signage).",
                    "Performance characteristics include: ability to be scored/cut/folded, moulding behaviour, transparency/translucency, flexibility, and whether it can be recycled or biodegrade.",
                    "Examples and uses (as per the spec): foam board (model making), fluted polypropylene (signs and box construction), translucent polypropylene sheets (packaging), styrofoam (modelling and formers)."
                ],
                "pitfalls": [
                    "Saying ‘plastic is waterproof’ without noting not all plastics are food-safe or UV-stable.",
                    "Mixing up transparent and translucent.",
                    "Ignoring end-of-life: some sheets are hard to recycle if mixed, dirty or laminated."
                ],
                "exam_application": {
                    "prompt": "Choose a polymer sheet for a semi-transparent package insert and justify your choice.",
                    "what_to_do": [
                        "Pick a suitable sheet (eg translucent polypropylene).",
                        "Explain: translucency, toughness, fold/crease behaviour, moisture resistance.",
                        "Mention manufacture (die cutting/creasing) and end-of-life."
                    ]
                }
            },
            {
                "title": "Woods and wood products: types + typical uses",
                "definitions": [
                    ("Softwood", "Usually from conifers; often easier to work and cheaper."),
                    ("Hardwood", "Usually from deciduous trees; often denser and more durable."),
                    ("MDF", "Medium density fibreboard: smooth, stable sheet made from wood fibres and resin."),
                    ("Veneer", "Thin layer of real wood glued to a substrate for appearance."),
                    ("Melamine formaldehyde laminate (MF)", "Hard decorative surface bonded to a board (wear resistant).")
                ],
                "explanations": [
                    "Woods are chosen for appearance, strength-to-weight, workability, and ‘warm’ feel. Manufactured boards are chosen for consistent size, stability and cost.",
                    "You should be familiar with these examples (spec list):\n"
                    "- Softwoods: pine, spruce, Douglas fir, redwood, cedar, larch\n"
                    "- Hardwoods: oak, ash, mahogany, teak, birch, beech\n"
                    "- Manufactured boards: plywood (including marine, aeroply, flexible plywood), chipboard, MDF\n"
                    "- Veneers and melamine formaldehyde laminates.",
                    "In exam reasoning, link the choice to stability (warping), surface quality (paint/laminate), and joining method (screws, dowels, KD fittings)."
                ],
                "pitfalls": [
                    "Calling MDF a natural timber (it is manufactured).",
                    "Ignoring wood movement (solid wood can expand/contract with moisture).",
                    "Choosing the wrong board for moisture (standard MDF swells if wet; marine plywood handles moisture better)."
                ],
                "exam_application": {
                    "prompt": "Select a wood product for a flat-pack shelf unit and justify.",
                    "what_to_do": [
                        "Choose a board (eg MDF with MF laminate, or plywood depending on context).",
                        "Explain: consistent thickness, good finish, suitable for KD fittings/screws.",
                        "Mention cost and batch production."
                    ]
                }
            },
            {
                "title": "Metals: stock forms, performance characteristics, and examples",
                "definitions": [
                    ("Stock form", "The standard shapes materials come in before processing (sheet, bar, tube, beam)."),
                    ("Alloying", "Mixing metals to improve properties (eg corrosion resistance, strength)."),
                    ("Melting point", "Temperature where solid becomes liquid; affects casting and heat processes."),
                    ("Coating", "Surface layer added for protection or appearance (paint, plating, anodising).")
                ],
                "explanations": [
                    "Metals are used when you need high strength, toughness, heat resistance, and good joining options.",
                    "Stock forms you should recognise include sheet, plate, bar (flat/round/square), tube (circular/hexagonal), and structural sections (H-beam, I-beam).",
                    "Key performance characteristics to describe include: malleability, toughness, hardness, density, corrosion resistance, thermal conductivity, electrical conductivity, melting points, ability to be alloyed, ability to be joined with heat processes, and ability to take coatings/finishes.",
                    "Metals you must be familiar with (spec list):\n"
                    "- Ferrous: low carbon steel, stainless steel, high speed steel (HSS), medium carbon steel, cast iron\n"
                    "- Non-ferrous: aluminium, copper, zinc, silver, gold, titanium, tin\n"
                    "- Ferrous alloys: stainless steel, die steel (tool steel)\n"
                    "- Non-ferrous alloys: bronze, brass, duralumin, pewter."
                ],
                "pitfalls": [
                    "Forgetting corrosion protection for ferrous metals (paint, plating, galvanising, stainless choice).",
                    "Assuming ‘strongest’ is always best (weight, cost, machining and corrosion also matter).",
                    "Ignoring that different stock forms change cost and waste (machining from solid creates waste)."
                ],
                "exam_application": {
                    "prompt": "A bracket must be strong, low cost, and produced in batch quantities. Choose a metal and a process.",
                    "what_to_do": [
                        "Select a suitable metal (eg low carbon steel) and justify (strength, cost, availability).",
                        "Select a process (press forming + punching; or laser cut + bend) and justify.",
                        "Mention corrosion protection (powder coat/galvanising) if needed."
                    ]
                }
            },
            {
                "title": "Polymers: stock forms, performance characteristics, and examples",
                "definitions": [
                    ("Insulation", "Not letting heat/electricity pass through easily."),
                    ("UV resistance", "How well a polymer survives sunlight without cracking/fading."),
                    ("Chemical resistance", "How well it survives oils/cleaners/solvents."),
                    ("Additives", "Extra ingredients to change performance (UV stabilisers, plasticisers).")
                ],
                "explanations": [
                    "Polymers are used for low weight, corrosion resistance, insulation, colourability, and fast mass production.",
                    "Stock forms include sheet, film, granules, rod/extruded forms, and foam.",
                    "Performance characteristics include: toughness, elasticity, insulation (thermal and electrical), UV resistance, ability to be moulded, resistance to chemicals/liquids, melting points, suitability for food packaging, biodegradability, and recyclability.",
                    "Polymers you must know (spec list):\n"
                    "- Thermoplastics: LDPE, HDPE, PP, HIPS, ABS, PMMA, nylon, rigid and flexible PVC, PET\n"
                    "- Thermosets: urea formaldehyde (UF), melamine formaldehyde (MF), polyester resin, epoxy resin.\n"
                    "The exam often wants you to link polymer choice to process (injection moulding, extrusion, blow moulding) and end-of-life (recycling codes, separation, contamination)."
                ],
                "pitfalls": [
                    "Saying ‘plastic is recyclable’ with no detail (some are, but collection/sorting matters).",
                    "Ignoring UV exposure (outdoor plastics often need UV stabilisers).",
                    "Not linking polymer choice to the shaping process and scale."
                ],
                "exam_application": {
                    "prompt": "A reusable food container needs to be tough, food safe, and suitable for injection moulding. Choose a polymer and justify.",
                    "what_to_do": [
                        "Choose PP or HDPE (common food-safe thermoplastics).",
                        "Explain toughness, chemical resistance, suitable melting behaviour for moulding.",
                        "Mention cleaning, durability, and recyclability."
                    ]
                }
            },
            {
                "title": "Elastomers: why they are used + examples",
                "definitions": [
                    ("Elastomer", "A polymer that can stretch and return to its original shape."),
                    ("Self-finishing", "A material that comes out of the process with a good surface finish (little finishing needed)."),
                    ("Durometer", "A measure of hardness for rubbers (useful word if asked to compare).")
                ],
                "explanations": [
                    "Elastomers are used for grip, sealing, shock absorption and comfort. The key behaviour is stretching and returning to shape.",
                    "Performance points to mention include: ability to stretch and return, texture (grip), and self-finishing surfaces.",
                    "Examples you should know (spec list): natural rubber, neoprene, silicone, and thermoplastic elastomer (TPE)."
                ],
                "pitfalls": [
                    "Calling any soft plastic an elastomer (some soft plastics are just plasticised PVC, not true elastomers).",
                    "Forgetting temperature and chemical limits (some rubbers swell in oils or go hard in cold)."
                ],
                "exam_application": {
                    "prompt": "A product needs a non-slip overmoulded grip. Explain why an elastomer is suitable and name one example material.",
                    "what_to_do": [
                        "Explain grip texture + flexibility + shock absorption.",
                        "Name a suitable elastomer (eg TPE or silicone) and mention overmoulding compatibility."
                    ]
                }
            },
            {
                "title": "Biodegradable polymers: suitability + how they degrade",
                "definitions": [
                    ("Biodegradable polymer", "A polymer designed to break down by sunlight, water or microorganisms (in the right conditions)."),
                    ("Enzymes", "Biological catalysts in soil/microbes that speed breakdown."),
                    ("Compostable", "Breaks down in controlled composting conditions (not always the same as biodegradable).")
                ],
                "explanations": [
                    "Biodegradable polymers are used to reduce long-term waste for certain products, especially short-life packaging.",
                    "The spec focus is: they can be moulded into 3D products or film, and they can degrade through UV (sunlight), water, or enzymes in soil. You should understand that degradation depends on conditions (heat, moisture, oxygen, microbes).",
                    "Examples you should know (spec list): corn starch polymers, potatopak, biopol (bio-batch additive), PLA, PHA, and some water-soluble lactide/glycolide materials (Lactel and ecofilm)."
                ],
                "pitfalls": [
                    "Assuming biodegradable means ‘disappears anywhere’ (it often needs specific conditions).",
                    "Using biodegradable polymers for long-life products where you actually want durability.",
                    "Not mentioning food-contact/safety requirements when relevant."
                ],
                "exam_application": {
                    "prompt": "A takeaway company wants film packaging that reduces environmental impact. Suggest a biodegradable polymer and explain how it degrades.",
                    "what_to_do": [
                        "Name PLA or starch-based film (depending on scenario).",
                        "Explain: can be made into film; breaks down via moisture/heat/microbes or UV (as appropriate).",
                        "Mention correct disposal route (eg composting/industrial processing if given)."
                    ]
                }
            },
            {
                "title": "Composites: what they are + why they outperform single materials",
                "definitions": [
                    ("Composite", "A material made by combining two or more materials to get better properties."),
                    ("Matrix", "The ‘binder’ part that holds the reinforcement (often resin or cement)."),
                    ("Reinforcement", "Fibres/particles that carry load and improve strength/stiffness."),
                    ("Laminating (layup)", "Layering reinforcement with resin to build thickness and shape.")
                ],
                "explanations": [
                    "Composites are chosen because they can be lighter for the same strength, stiffer for the same weight, or tougher and more durable than a single material.",
                    "Spec points to mention include: they can be moulded into a variety of 3D forms; they can enhance physical/mechanical properties; sometimes they are easier to manufacture than traditional solutions; and they can improve product performance.",
                    "Examples you should know (spec list): CFRP, GRP, tungsten carbide, aluminium composite board, concrete (including reinforced concrete), fibre cement, and engineered wood such as glulam."
                ],
                "pitfalls": [
                    "Saying composites are always recyclable (many fibre-resin composites are difficult to recycle).",
                    "Not explaining what is ‘combined’ (matrix + reinforcement).",
                    "Ignoring directionality (fibre orientation can change strength/stiffness)."
                ],
                "exam_application": {
                    "prompt": "Explain why CFRP might be used instead of aluminium for a high-performance product.",
                    "what_to_do": [
                        "Compare strength-to-weight/stiffness-to-weight.",
                        "Mention performance benefits (fatigue resistance, design freedom).",
                        "Mention trade-offs: cost, manufacture complexity, recycling."
                    ]
                }
            },
            {
                "title": "Smart materials: responding to stimuli",
                "definitions": [
                    ("Smart material", "A material that changes property/shape in response to a stimulus."),
                    ("Stimulus", "The trigger: temperature, light, pressure/force, electricity."),
                    ("Shape memory alloy (SMA)", "Metal that returns to a preset shape when heated (eg Nitinol).")
                ],
                "explanations": [
                    "Smart materials are used when you want the material to ‘do something’ automatically, like change colour, glow, or move, without complex mechanisms.",
                    "The spec wants you to link suitability to the stimulus response: changes in temperature, light levels, or pressure/force.",
                    "Examples you should know (spec list): SMA (Nitinol), thermochromic pigment, phosphorescent pigment, photochromic pigment, electroluminescent wire, piezoelectric material."
                ],
                "pitfalls": [
                    "Describing a smart material but not naming the stimulus (what triggers it).",
                    "Forgetting limitations (temperature range, lifetime, cost).",
                    "Saying ‘phosphorescent’ is the same as ‘reflective’ (it glows after being charged by light)."
                ],
                "exam_application": {
                    "prompt": "Suggest a smart material feature for a safety product and explain why it is suitable.",
                    "what_to_do": [
                        "Pick one (eg photochromic for light change, phosphorescent for glow in dark).",
                        "Explain stimulus and response.",
                        "Link to user benefit (visibility, warning, feedback)."
                    ]
                }
            },
            {
                "title": "Modern materials: advanced options and why you would choose them",
                "definitions": [
                    ("Modern material", "A newer/advanced material developed for high performance or new applications."),
                    ("Kevlar", "High strength fibre with excellent impact resistance (often used as reinforcement)."),
                    ("Polymorph", "Low-melting thermoplastic that can be reshaped by heating in hot water.")
                ],
                "explanations": [
                    "Modern materials are often chosen for performance advantages (strength-to-weight, modelling speed, special forming behaviour).",
                    "Examples you should know (spec list): Kevlar, precious metal clay (PMC), high density modelling foam, and Polymorph.",
                    "In exam answers, give the reason in plain terms: ‘Kevlar reinforcement helps absorb impacts while staying light’, or ‘modelling foam is fast to shape accurately for prototypes’."
                ],
                "pitfalls": [
                    "Naming a modern material without explaining the advantage for that product.",
                    "Forgetting cost and manufacturing limitations.",
                    "Using a modern material when a standard one would do (exams often want ‘appropriate’ not ‘fancy’)."
                ],
                "exam_application": {
                    "prompt": "A company needs a fast prototype that can be reshaped multiple times. Suggest a modern material and explain.",
                    "what_to_do": [
                        "Choose Polymorph or high density modelling foam depending on brief.",
                        "Explain quick forming and ability to rework.",
                        "Link to prototyping speed and cost."
                    ]
                }
            },
        ]
    },

    {
        "section": "3.1.3",
        "title": "Enhancement of materials",
        "topics": [
            {
                "title": "Polymer enhancement (additives)",
                "definitions": [
                    ("Additive", "Something mixed into a polymer to change its performance."),
                    ("UV stabiliser", "Additive that slows damage from sunlight (stops fading/cracking)."),
                    ("Bio-batch additive", "Additive designed to increase biodegradability of a polymer.")
                ],
                "explanations": [
                    "Polymers can be ‘tuned’ using additives. This is common in real products because it is cheaper than changing the whole material.",
                    "UV stabilisers are used for outdoor plastics (eg patio furniture) to prevent sunlight from making them brittle or discoloured.",
                    "Bio-batch additives are used to encourage biodegradability for some products like certain packaging/carrier bags. Always link it to the product lifespan: you want it to last long enough for use, but not forever as waste."
                ],
                "pitfalls": [
                    "Saying additives always make a product ‘more sustainable’ (some additives can make recycling harder).",
                    "Not stating the reason for the additive (eg UV stabiliser because of outdoor use)."
                ],
                "exam_application": {
                    "prompt": "A plastic garden product fails after one summer outdoors. Suggest an enhancement method.",
                    "what_to_do": [
                        "Recommend UV stabilisers (or UV-resistant polymer) and explain the sunlight damage mechanism.",
                        "Link to longer life and reduced waste."
                    ]
                }
            },
            {
                "title": "Wood enhancement (resins, lamination, preservatives)",
                "definitions": [
                    ("Lamination", "Bonding layers together to improve strength/stability."),
                    ("Resin", "A polymer ‘binder’ used to strengthen or protect wood products."),
                    ("Preservative", "Chemical that prevents decay, rot or insect attack.")
                ],
                "explanations": [
                    "Natural timber can be combined with resins and laminated to increase strength and stability (less warping, better consistency).",
                    "Preservatives and protective finishes/coatings are used when wood is exposed to moisture, UV and outdoor conditions. They improve lifespan and reduce maintenance.",
                    "A good exam answer links enhancement to the failure you are preventing: rot/decay, moisture swelling, surface wear, or colour fade."
                ],
                "pitfalls": [
                    "Not linking the enhancement to the environment (indoor vs outdoor).",
                    "Confusing decorative finishes with preservatives (a stain may colour wood but not stop rot)."
                ],
                "exam_application": {
                    "prompt": "A wooden outdoor product needs improved durability. Suggest enhancements.",
                    "what_to_do": [
                        "Use preservatives/pressure treatment and a suitable finish (oil/varnish).",
                        "Explain how it prevents decay and water ingress."
                    ]
                }
            },
            {
                "title": "Metal enhancement (heat treatment)",
                "definitions": [
                    ("Heat treatment", "Heating and cooling metal in a controlled way to change properties."),
                    ("Case hardening", "Hardens the surface while keeping a tougher core (wear resistance)."),
                    ("Hardening and tempering", "Hardening increases hardness; tempering reduces brittleness and improves toughness.")
                ],
                "explanations": [
                    "Heat treatment is used to make metals better for their job without changing the shape.",
                    "Case hardening is used for parts that need a hard-wearing surface (gears, shafts) but must not be brittle through the whole thickness.",
                    "Hardening and tempering is a common pair: harden to get the hardness you need, then temper to stop it snapping in use."
                ],
                "pitfalls": [
                    "Saying ‘harder is always better’ (too hard can be brittle).",
                    "Not linking the process to the product requirement (wear vs impact vs fatigue)."
                ],
                "exam_application": {
                    "prompt": "A steel component wears out quickly. Suggest an enhancement method.",
                    "what_to_do": [
                        "Recommend case hardening for wear resistance (or harden and temper).",
                        "Explain surface hardness vs core toughness."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.4",
        "title": "Forming, redistribution and addition processes",
        "topics": [
            {
                "title": "Paper and board forming processes (packaging and nets)",
                "definitions": [
                    ("Die cutting", "Cutting shapes using a custom die (fast for batch/mass)."),
                    ("Creasing", "Making a fold line so board folds cleanly."),
                    ("Laser cutting", "Using a laser for accurate cutting (good for prototypes/small batches)."),
                    ("Bending", "Forming paper/board into shape by folding/curving (often with creases).")
                ],
                "explanations": [
                    "Paper and board are shaped mainly with cutting and folding operations. You often start with a 2D net, then cut and crease it, then fold into 3D packaging.",
                    "Die cutting is efficient for high volumes (once the die is made). Laser cutting is slower but flexible (no die cost), so it suits prototypes and small runs."
                ],
                "pitfalls": [
                    "Not mentioning that dies have high setup cost (tooling) but low unit cost.",
                    "Forgetting that thick board needs creasing before folding.",
                    "Choosing laser cutting for mass production without justifying time/cost."
                ],
                "exam_application": {
                    "prompt": "A company needs 50,000 packaging sleeves. Choose a forming method and justify.",
                    "what_to_do": [
                        "Choose die cutting + creasing (tooling cost is worth it at volume).",
                        "Mention speed, repeatability, accuracy, and print registration."
                    ]
                }
            },
            {
                "title": "Polymer forming processes (turning sheet/granules into 3D)",
                "definitions": [
                    ("Vacuum forming", "Heated sheet pulled over mould by vacuum."),
                    ("Thermoforming", "General term for forming heated sheet (vacuum forming is one type)."),
                    ("Calendaring", "Rolling hot polymer into sheet/film of controlled thickness."),
                    ("Injection moulding", "Molten polymer injected into a mould; excellent for mass production."),
                    ("Blow moulding", "Forming hollow products by inflating a hot tube inside a mould (bottles)."),
                    ("Rotational moulding", "Powder melts inside rotating mould; hollow items; low pressure."),
                    ("Extrusion", "Forcing molten polymer through a die to make continuous shapes (tubes, profiles)."),
                    ("Compression moulding", "Polymer pressed into a heated mould; common for some thermosets."),
                    ("Line bending", "Heating a line in sheet plastic so it bends sharply and neatly."),
                    ("Laminating (layup)", "Building layers (often in composites) to form a shape.")
                ],
                "explanations": [
                    "The ‘best’ polymer process depends on product shape, material type, and production volume.",
                    "Injection moulding has high tooling cost but very low unit cost and excellent repeatability - ideal for mass production of detailed 3D parts.",
                    "Vacuum forming/thermoforming is great for thin-walled packaging and trays. Extrusion is for long continuous products. Blow moulding is for hollow containers. Rotational moulding suits large hollow parts with even wall thickness but slower cycle times."
                ],
                "pitfalls": [
                    "Not linking process to volume (eg choosing injection moulding for a one-off prototype).",
                    "Mixing up blow moulding (hollow) vs injection moulding (solid parts).",
                    "Forgetting wall thickness limits and draft angles for mould release."
                ],
                "exam_application": {
                    "prompt": "Choose a process to make 1 million plastic bottles and justify.",
                    "what_to_do": [
                        "Choose blow moulding (often with preforms depending on polymer).",
                        "Explain hollow shape, speed, low unit cost at high volume."
                    ]
                }
            },
            {
                "title": "Metal shaping, joining, and wasting processes",
                "definitions": [
                    ("Press forming", "Using a press/die to form sheet metal into shapes."),
                    ("Spinning", "Forming sheet over a rotating former (often circular parts)."),
                    ("Deep drawing", "Pressing sheet into deep cup shapes without tearing."),
                    ("Forging / drop forging", "Shaping hot metal with compressive forces; improves grain structure."),
                    ("Rolling", "Reducing thickness / shaping using rollers."),
                    ("Casting", "Pouring molten metal into a mould to set into shape."),
                    ("MIG/TIG welding", "Common metal welding methods (different control and typical uses)."),
                    ("Soldering / brazing", "Joining using a filler metal; brazing is higher temp/stronger than soft solder."),
                    ("Riveting", "Permanent mechanical joining using rivets."),
                    ("Milling / turning", "Machining (wasting) processes to remove material to shape.")
                ],
                "explanations": [
                    "Metal processes are chosen based on stock form and desired strength/finish. Sheet metal parts often use press forming, bending and punching. Round ‘dish’ parts can use spinning. Deep drawing forms cups and cans.",
                    "Casting is good for complex shapes; different methods (sand, die, investment, low-temp casting like pewter) trade off cost, accuracy and surface finish.",
                    "Joining can be by welding (MIG, TIG, spot, oxy-acetylene), by soldering/brazing, or mechanical methods like rivets and bolts. Wasting processes (milling, turning, flame/plasma/laser cutting, punching/stamping) remove material and can create waste."
                ],
                "pitfalls": [
                    "Not identifying whether the part starts as sheet, bar or cast stock (this changes the best method).",
                    "Ignoring heat effects (welding can distort thin sheet).",
                    "Choosing machining from solid for a simple sheet part (creates unnecessary waste/cost)."
                ],
                "exam_application": {
                    "prompt": "A steel bracket needs to be made in batch quantities from sheet metal. Pick suitable processes.",
                    "what_to_do": [
                        "Suggest laser cutting or punching/stamping then bending/press forming.",
                        "Mention repeatability, speed, low waste, and finishing (powder coat/galvanise)."
                    ]
                }
            },
            {
                "title": "Wood processes: joining and forming into 3D",
                "definitions": [
                    ("Traditional joint", "Wood-to-wood joint formed by shaping both parts (eg mortise and tenon)."),
                    ("Knock-down (KD) fitting", "Hardware that allows flat-pack assembly/disassembly."),
                    ("Steam bending", "Softening wood with steam then bending around a former."),
                    ("Laminating", "Layering wood veneers/strips with glue to make a strong curved form.")
                ],
                "explanations": [
                    "Wood products often combine joints and fixings. Traditional joints (dovetail, comb, housing, half-lap, dowel, mortise and tenon) are used where strength and appearance matter.",
                    "Component jointing (KD fittings, wood screws, nuts/bolts, coach bolts) is used when you need fast assembly, repair, or flat-pack delivery.",
                    "Forming wood into 3D shapes can be done with laminating (strong curved parts), steam bending (natural curves), and machine processes such as turning (between centres, using chuck/faceplate), milling and routing (slots/holes)."
                ],
                "pitfalls": [
                    "Choosing the prettiest joint without considering scale of production and time.",
                    "Ignoring wood movement and grain direction when bending/laminating.",
                    "Forgetting that KD fittings need appropriate board thickness and edge strength."
                ],
                "exam_application": {
                    "prompt": "A chair needs a curved wooden backrest with consistent shape. Suggest a process.",
                    "what_to_do": [
                        "Suggest laminating thin layers around a former (repeatable and strong).",
                        "Explain stability compared with steam bending for mass/batch production."
                    ]
                }
            },
            {
                "title": "Adhesives and fixings (what to use and why)",
                "definitions": [
                    ("PVA", "Water-based wood glue; strong for porous materials, needs clamp time."),
                    ("Contact adhesive", "Instant bond when two coated surfaces touch; good for laminates."),
                    ("UV hardening adhesive", "Cures quickly when exposed to UV light; fast production."),
                    ("Solvent cement", "Dissolves surfaces (eg acrylic) to fuse them (eg Tensol/acrylic cement)."),
                    ("Epoxy resin", "Two-part adhesive; strong, gap-filling, bonds many materials.")
                ],
                "explanations": [
                    "Adhesive choice depends on materials, joint type, time available, and durability needs.",
                    "PVA is excellent for wood and paper-based materials. Contact adhesives are common for laminates because they bond large areas quickly. UV adhesives are used when speed matters and parts can be exposed to UV light. Solvent cements are used for plastics like acrylic to create an almost ‘welded’ joint. Epoxy is used when you need high strength and gap filling across different materials."
                ],
                "pitfalls": [
                    "Using PVA on non-porous plastics or metals (won’t bond well).",
                    "Not preparing surfaces (grease/dust ruins joints).",
                    "Forgetting safety: ventilation/PPE for solvents and some resins."
                ],
                "exam_application": {
                    "prompt": "A clear acrylic display needs a strong, neat joint. Choose an adhesive and justify.",
                    "what_to_do": [
                        "Use solvent cement (acrylic cement/Tensol) for clear, neat joints.",
                        "Explain ‘fusing’ action and clean finish."
                    ]
                }
            },
            {
                "title": "Jigs and fixtures (accuracy and repeatability)",
                "definitions": [
                    ("Jig", "A tool that guides the cutting/drilling tool to the right place."),
                    ("Fixture", "A tool that holds the workpiece in the correct position."),
                    ("Template", "A pattern used to mark out or guide shaping repeatedly.")
                ],
                "explanations": [
                    "Jigs and fixtures are used to make manufacture accurate and repeatable. They reduce human error, speed up production, and improve consistency.",
                    "In exam questions, you usually describe how a jig/fixture would locate the part (datum points), clamp it, and control the tool path (drill guide, router template)."
                ],
                "pitfalls": [
                    "Saying ‘use a jig’ without explaining what it does (holds, locates, guides).",
                    "Not mentioning accuracy/repeatability as the main benefit.",
                    "Forgetting safety: jigs often reduce risk by keeping hands away from tools."
                ],
                "exam_application": {
                    "prompt": "A workshop needs 200 identical holes drilled in the same position. Explain how a jig helps.",
                    "what_to_do": [
                        "Describe a drill jig with a guide bush and locating faces.",
                        "Explain repeatability, speed, and reduced measurement errors."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.5",
        "title": "The use of finishes",
        "topics": [
            {
                "title": "Paper and board finishes + printing processes",
                "definitions": [
                    ("Lamination", "Bonding a thin film to the surface for protection/appearance."),
                    ("Embossing / debossing", "Raising / pressing in a design for texture and premium feel."),
                    ("Varnish (UV/spot)", "Clear coating; UV cures fast; spot varnish highlights areas."),
                    ("Screen printing", "Ink pushed through a mesh stencil; great for bold graphics."),
                    ("Flexographic printing", "Fast printing for packaging using flexible plates."),
                    ("Offset lithography", "High-quality printing using plates; common for large runs."),
                    ("Digital printing", "No plates; fast setup; ideal for small runs or variable data.")
                ],
                "explanations": [
                    "Finishes change appearance and function (water resistance, scuff resistance, premium feel).",
                    "Lamination improves durability and moisture resistance but can reduce recyclability. Embossing/debossing adds texture without extra ink. Varnishes protect prints and can create contrast (spot varnish).",
                    "Printing process choice depends on volume, quality, and cost: digital for low volume/variable data, offset for high quality large runs, flexo for packaging at high speed, and screen printing for strong colours and special effects."
                ],
                "pitfalls": [
                    "Picking a finish without considering recyclability (laminates can complicate recycling).",
                    "Saying ‘digital is best’ for huge runs (unit cost may be higher than offset/flexo).",
                    "Ignoring abrasion/moisture requirements for packaging."
                ],
                "exam_application": {
                    "prompt": "Choose a print and finish method for a premium gift box produced in medium volume.",
                    "what_to_do": [
                        "Suggest offset litho for quality + embossing/spot UV varnish for premium feel.",
                        "Explain trade-off: cost vs appearance and durability."
                    ]
                }
            },
            {
                "title": "Polymer finishes (self-finishing, coatings and effects)",
                "definitions": [
                    ("Self-finishing polymer", "Comes out of the mould with a good surface (minimal finishing)."),
                    ("Acrylic spray paint", "Common coating to add colour/protection to plastics."),
                    ("Gel coat", "Protective surface layer used with GRP (improves finish and durability)."),
                    ("Smart pigment", "Pigment that changes (thermochromic) or glows (phosphorescent)."),
                    ("Thermoplastic elastomer (TPE) finish", "Soft-touch/grip finish often achieved by overmoulding.")
                ],
                "explanations": [
                    "Many plastics are designed to be self-finishing: the mould surface texture creates the final look.",
                    "Where extra finish is needed, acrylic spray paints can add colour. GRP often uses gel coats for a smooth, durable outer layer.",
                    "Smart pigments can add user feedback or visual effects, like temperature change colour (thermochromic) or glow (phosphorescent)."
                ],
                "pitfalls": [
                    "Painting a polymer without proper surface prep (paint can peel).",
                    "Choosing an effect pigment without linking it to user need (marks come from reasoning)."
                ],
                "exam_application": {
                    "prompt": "A sports product needs a durable coloured polymer surface. Suggest a finishing approach.",
                    "what_to_do": [
                        "Use moulded-in colour/self-finishing where possible; otherwise choose appropriate coating (eg acrylic spray) and explain prep."
                    ]
                }
            },
            {
                "title": "Metal finishes (appearance and corrosion protection)",
                "definitions": [
                    ("Electroplating", "Using electricity to coat metal with another metal (appearance/corrosion)."),
                    ("Powder coating", "Dry powder baked on; tough protective finish."),
                    ("Galvanising", "Zinc coating on steel to prevent rust."),
                    ("Anodising", "Electrochemical process forming a protective oxide layer on aluminium."),
                    ("Cathodic protection", "Protecting metal from corrosion by making it the cathode (sacrificial anodes).")
                ],
                "explanations": [
                    "Metal finishes are chosen to prevent corrosion, improve wear resistance, and achieve a desired look.",
                    "Paints (cellulose/acrylic) are simple coatings. Powder coating is durable for outdoor/industrial use. Electroplating gives a metal look and can improve corrosion resistance. Galvanising is a go-to for steel outdoors. Anodising is common for aluminium for a tough, coloured finish.",
                    "In exam answers, always link the finish to the environment (water, salt, outdoor) and expected lifespan."
                ],
                "pitfalls": [
                    "Not stating why corrosion is a risk (environment).",
                    "Mixing ‘plating’ and ‘anodising’ (anodising is for aluminium oxide layer).",
                    "Choosing decorative coating only when protection is needed."
                ],
                "exam_application": {
                    "prompt": "A steel outdoor product is rusting. Recommend a finish and justify.",
                    "what_to_do": [
                        "Recommend galvanising or powder coating (depending on design).",
                        "Explain barrier/sacrificial protection and durability."
                    ]
                }
            },
            {
                "title": "Wood finishes (appearance and decay prevention)",
                "definitions": [
                    ("Varnish", "Clear protective film (polyurethane/acrylic) that seals wood."),
                    ("Stain / colour wash", "Adds colour while showing grain."),
                    ("Wax / oils (Danish/teak)", "Soaks in and protects; gives natural feel; needs reapplication."),
                    ("Pressure treatment", "Forcing preservatives into wood under pressure for deep protection.")
                ],
                "explanations": [
                    "Wood finishes either form a protective layer (varnish/paint) or soak in (oils/waxes).",
                    "For outdoor use, you often need preservatives and a finish to limit moisture ingress and UV damage. Pressure treatment is used when deep, long-lasting protection is needed.",
                    "Spec examples include polyurethane varnish, acrylic varnish, water-based paints, stains, colour wash, wax finishes, Danish oil, teak oil, and pressure treating with chemical preservatives."
                ],
                "pitfalls": [
                    "Using indoor finishes outdoors (they break down quickly).",
                    "Assuming a stain is a preservative (it mainly colours).",
                    "Not mentioning maintenance (oils often need reapplication)."
                ],
                "exam_application": {
                    "prompt": "Choose a finish for an outdoor wooden bench and explain your choice.",
                    "what_to_do": [
                        "Use pressure treated timber + an outdoor-grade finish (oil/varnish).",
                        "Explain moisture/UV protection and maintenance."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.6",
        "title": "Modern industrial and commercial practice",
        "topics": [
            {
                "title": "Scales of production (one-off to mass)",
                "definitions": [
                    ("One-off / bespoke", "Made once or in very small quantities; highly customised."),
                    ("Batch production", "Made in groups; equipment set up for a run, then changed."),
                    ("Mass / line production", "Very high volume, often on production lines; low unit cost.")
                ],
                "explanations": [
                    "Scale affects everything: material choice, process choice, tooling cost, and how you control quality.",
                    "One-off work suits flexible processes (manual methods, CNC/laser cutting) because tooling must be cheap. Mass production can justify expensive tooling (moulds, dies) because cost is spread over many units.",
                    "In exam questions, you usually pick a scale then justify it based on volume, variety, cost, and speed."
                ],
                "pitfalls": [
                    "Saying ‘mass production is cheapest’ without noting high setup/tooling costs.",
                    "Choosing a process that does not match the proposed volume."
                ],
                "exam_application": {
                    "prompt": "A product is expected to sell 5,000 units a year. Which scale of production is suitable and why?",
                    "what_to_do": [
                        "Likely batch production.",
                        "Explain balance of setup cost vs unit cost, flexibility for updates, manageable tooling."
                    ]
                }
            },
            {
                "title": "Cost, waste reduction, and efficient manufacture (including JIT)",
                "definitions": [
                    ("Material economy", "Using materials efficiently with minimal waste."),
                    ("Waste", "Offcuts, rejects, excess processing, energy loss."),
                    ("Just In Time (JIT)", "Materials/components arrive when needed to reduce storage costs.")
                ],
                "explanations": [
                    "The spec wants you to understand the link between material cost/form, manufacturing process, and scale.",
                    "Designs can reduce waste by using standard sizes, nesting parts on sheets, reducing machining, and reducing part count. Processes that increase accuracy reduce rejects.",
                    "JIT reduces storage cost and can improve efficiency, but it relies on reliable suppliers and good planning."
                ],
                "pitfalls": [
                    "Ignoring that reducing waste can also reduce labour time and rejects.",
                    "Saying JIT is always better (supply delays can stop the whole line)."
                ],
                "exam_application": {
                    "prompt": "Explain two ways to reduce waste when manufacturing a sheet-based product.",
                    "what_to_do": [
                        "Mention nesting/CAD layout, standard sizes, fewer operations, consistent tolerances.",
                        "Link to cost and sustainability."
                    ]
                }
            },
            {
                "title": "Computer systems in manufacture + manufacturing systems",
                "definitions": [
                    ("Unit Production Systems (UPS)", "Workstations/tools that can be reconfigured for different products."),
                    ("Quick Response Manufacturing (QRM)", "Focus on reducing lead time to respond quickly to demand."),
                    ("Vertical in-house production", "Company makes more parts internally rather than outsourcing."),
                    ("Modular/cell production", "Work organised into cells/modules to improve flow and flexibility."),
                    ("Flexible Manufacturing System (FMS)", "Automated system that can switch between products with minimal downtime.")
                ],
                "explanations": [
                    "Computer systems plan and control manufacturing: scheduling machines/people, managing stock, coordinating suppliers, and reducing waste.",
                    "UPS and cell production help factories adapt when products change. QRM focuses on speed (short lead times). Vertical in-house production can improve quality control and reliability but requires investment.",
                    "FMS uses automation and computer control to produce different products efficiently - useful when you need variety without huge downtime."
                ],
                "pitfalls": [
                    "Listing system names without explaining the benefit (speed, flexibility, cost).",
                    "Forgetting the trade-off (automation costs money; in-house production needs investment)."
                ],
                "exam_application": {
                    "prompt": "A company frequently updates products and wants fast changeovers. Which manufacturing system helps and why?",
                    "what_to_do": [
                        "Choose modular/cell production, UPS, or FMS depending on scenario.",
                        "Explain reduced setup time and improved flexibility."
                    ]
                }
            },
            {
                "title": "Sub-assembly (building parts separately)",
                "definitions": [
                    ("Sub-assembly", "Making a part of a product on a separate line before final assembly.")
                ],
                "explanations": [
                    "Sub-assembly improves efficiency because specialists can build modules in parallel, then final assembly becomes quicker and more consistent.",
                    "It also helps quality control: if a module fails, you replace the module rather than reworking the whole product."
                ],
                "pitfalls": [
                    "Not linking sub-assembly to speed, quality and consistency.",
                    "Ignoring logistics (stocking modules, version control)."
                ],
                "exam_application": {
                    "prompt": "Explain one advantage of sub-assembly in a mass-produced product.",
                    "what_to_do": [
                        "Mention parallel production, faster final assembly, easier quality checks, simpler repair."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.7",
        "title": "Digital design and manufacture",
        "topics": [
            {
                "title": "Computer Aided Design (CAD): what it’s for and why it matters",
                "definitions": [
                    ("CAD", "Software used to create 2D drawings and 3D models of designs."),
                    ("2D CAD", "Accurate drawings (dimensions, tolerances, layouts)."),
                    ("3D CAD", "Models for visualisation, assemblies, and sometimes simulation."),
                    ("Tolerance", "Allowed variation in a dimension to ensure fit and function."),
                    ("Datum", "A reference point/edge used to measure and locate features.")
                ],
                "explanations": [
                    "CAD improves accuracy, speed of edits, and communication. It can generate working drawings, assemblies, and visuals for clients/manufacturers.",
                    "Compared with manual drawing, CAD is faster to revise and easier to share, but it needs training, hardware, and good file management.",
                    "In industry, CAD links with CAM and data systems to reduce errors between design and manufacture."
                ],
                "pitfalls": [
                    "Saying CAD is always better (manual sketching can be faster for early ideas).",
                    "Not mentioning tolerances/datum points when discussing working drawings.",
                    "Forgetting that CAD files need version control (wrong file can cause manufacturing mistakes)."
                ],
                "exam_application": {
                    "prompt": "Give two advantages and one disadvantage of CAD compared to manual methods.",
                    "what_to_do": [
                        "Advantages: accuracy, quick edits, easy sharing, integration with CAM/simulation.",
                        "Disadvantage: training cost, software/hardware cost, risk of over-detailing early ideas."
                    ]
                }
            },
            {
                "title": "Digital manufacture (CAM/CNC): what processes do",
                "definitions": [
                    ("CAM", "Computer Aided Manufacture: using digital files to control machines."),
                    ("CNC", "Computer Numerical Control: machines following programmed tool paths."),
                    ("Laser cutting", "Cuts sheet material accurately; good for 2D profiles."),
                    ("Routing", "CNC cutter removes material (often wood/plastics) to shape."),
                    ("Milling/turning", "CNC machining processes for metals and plastics."),
                    ("Plotter cutting", "Knife cutting for thin sheet/film/vinyl.")
                ],
                "explanations": [
                    "CAM uses CAD data to make parts accurately and repeatedly. It reduces human marking errors and can create complex shapes.",
                    "Laser cutting is fast for sheet profiles. Routing is common for wood and plastics. Milling and turning are used for precision metal/plastic parts. Plotter cutting is used for vinyl graphics and thin films.",
                    "A strong exam answer links the process to tolerances, repeatability, and production scale."
                ],
                "pitfalls": [
                    "Calling every digital process ‘3D printing’ (not in this section’s list).",
                    "Not linking CAM to accuracy and repeatability.",
                    "Ignoring setup time and the need for correct tooling/feeds/speeds."
                ],
                "exam_application": {
                    "prompt": "A company needs accurate sheet components for a small batch. Choose a CAM process and justify.",
                    "what_to_do": [
                        "Laser cut or CNC routed depending on material.",
                        "Explain accuracy, repeatability, no tooling (vs dies), and speed for small batch."
                    ]
                }
            },
            {
                "title": "Simulation and analysis (CFD and FEA)",
                "definitions": [
                    ("Simulation", "Testing a design virtually before making it."),
                    ("CFD", "Computational Fluid Dynamics: simulating airflow/liquid flow."),
                    ("FEA", "Finite Element Analysis: simulating stresses/strain and deformation.")
                ],
                "explanations": [
                    "Simulation reduces prototype cost and speeds up development. CFD is used for products affected by airflow/fluids (fans, vehicles, vents). FEA is used to predict where a part will bend, crack or fail under load.",
                    "In exam questions, you explain what the tool tells you and how that improves design (reduce weight, improve strength, reduce drag, reduce overheating)."
                ],
                "pitfalls": [
                    "Naming CFD/FEA without saying what they actually analyse.",
                    "Claiming simulation ‘proves’ something (it predicts; still needs real testing)."
                ],
                "exam_application": {
                    "prompt": "Explain how FEA could help improve a bracket design.",
                    "what_to_do": [
                        "Describe stress map, identifying weak points, adding ribs/fillets, reducing weight where stress is low."
                    ]
                }
            },
            {
                "title": "Electronic data interchange and EPOS (data for manufacturing and marketing)",
                "definitions": [
                    ("EPOS", "Electronic Point Of Sale: digital checkout systems that collect sales data."),
                    ("Electronic data interchange", "Digital sharing of data between businesses (orders, stock, schedules)."),
                    ("Stock control", "Keeping correct inventory levels without over/under stocking.")
                ],
                "explanations": [
                    "EPOS and linked systems help businesses understand demand and manage supply. The spec expects you to explain uses such as: maintaining stock levels, capturing customer data (contact details), checking material availability, scheduling machines/people, and coordinating suppliers/customers.",
                    "In simple terms: good data helps you make the right amount at the right time, with less waste and fewer delays."
                ],
                "pitfalls": [
                    "Only talking about ‘marketing’ and forgetting manufacturing planning (scheduling/stock).",
                    "Not mentioning data protection/privacy when customer data is involved."
                ],
                "exam_application": {
                    "prompt": "Explain how EPOS data can reduce waste in manufacturing.",
                    "what_to_do": [
                        "Link accurate demand forecasting to making the right quantities, reducing overproduction and excess stock."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.8",
        "title": "The requirements for product design and development",
        "topics": [
            {
                "title": "Product development and improvement (learning from existing products)",
                "definitions": [
                    ("Fitness for purpose", "How well the product does its intended job."),
                    ("Specification criteria", "The requirements the product must meet."),
                    ("Critical assessment", "Judging strengths/weaknesses using evidence and user needs.")
                ],
                "explanations": [
                    "The spec wants you to use existing products as learning tools: analyse what works, what fails, and why.",
                    "This builds your ability to design to a specification, ensure fitness for purpose, and achieve accuracy of production.",
                    "When you critique products properly, you can justify changes that lead to improved or completely new designs. You should also consider aesthetics, ergonomics and anthropometrics (how it looks and how it fits people)."
                ],
                "pitfalls": [
                    "Saying ‘it’s good/bad’ with no evidence (needs reasons and user link).",
                    "Ignoring manufacturing accuracy (a product can be well-designed but poorly made).",
                    "Forgetting user ergonomics/anthropometrics."
                ],
                "exam_application": {
                    "prompt": "Analyse a product and suggest two improvements linked to user needs and manufacture.",
                    "what_to_do": [
                        "State the issue, the user impact, and a realistic manufacturing change.",
                        "Use precise terms: ergonomics, tolerance, material choice, process change."
                    ]
                }
            },
            {
                "title": "Inclusive design (designing for a wide range of users)",
                "definitions": [
                    ("Inclusive design", "Designing so a product can be used by as many people as possible."),
                    ("Accessibility", "How usable a product is for people with different abilities."),
                    ("Anthropometrics", "Human body measurements used to set sizes/clearances.")
                ],
                "explanations": [
                    "Inclusive design is about reducing barriers. That can mean larger grips, clearer labels, better contrast, tactile feedback, or simpler assembly.",
                    "The spec highlights users such as disabled people, children and the elderly. So your reasoning should mention strength, dexterity, vision, and safety."
                ],
                "pitfalls": [
                    "Saying ‘make it bigger’ without linking to a specific user need (grip strength, vision, dexterity).",
                    "Forgetting trade-offs (bigger buttons may affect aesthetics/space)."
                ],
                "exam_application": {
                    "prompt": "Explain two inclusive design features for a kitchen product.",
                    "what_to_do": [
                        "Feature + reason: high-contrast markings for vision; larger textured grip for weaker hands; clear feedback (clicks) for confidence."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.9",
        "title": "Health and safety",
        "topics": [
            {
                "title": "Safe working practices and legislation (workshops and manufacture)",
                "definitions": [
                    ("Health and Safety at Work Act (1974)", "UK law requiring safe working conditions and procedures."),
                    ("COSHH", "Control of Substances Hazardous to Health: rules for hazardous materials."),
                    ("Risk assessment", "Identifying hazards, assessing risk, and putting controls in place."),
                    ("Hazard", "Something that can cause harm."),
                    ("Risk", "How likely harm is and how serious it could be.")
                ],
                "explanations": [
                    "Paper 1 expects you to understand safe manufacture and how legislation influences it.",
                    "COSHH covers things like adhesives, solvents, finishes and dust - you must consider ventilation, PPE, storage, and correct disposal.",
                    "Risk assessment is a simple idea: spot hazards, decide how bad they could be, then reduce risk with controls (guards, PPE, training, procedures)."
                ],
                "pitfalls": [
                    "Giving generic safety statements without matching the hazard (eg ‘wear goggles’ for everything).",
                    "Forgetting fumes/dust and ventilation for finishes and solvents.",
                    "Confusing hazard and risk."
                ],
                "exam_application": {
                    "prompt": "Write a short risk assessment for using a solvent-based adhesive.",
                    "what_to_do": [
                        "Identify hazard (fumes/flammability/skin contact).",
                        "Controls (ventilation, gloves, no flames, storage, disposal).",
                        "Mention COSHH and safe working procedures."
                    ]
                }
            },
            {
                "title": "Consumer safety, standards and instructions",
                "definitions": [
                    ("Consumer legislation", "Laws to protect people who buy/use products."),
                    ("British Standards Institute (BSI)", "Organisation that sets standards for safety and quality."),
                    ("Safety warning", "Clear instruction about hazards to prevent misuse.")
                ],
                "explanations": [
                    "Products must be safe in normal use and reasonably foreseeable misuse. Standards (eg BSI) help define safety requirements.",
                    "The spec also expects you to describe advice to consumers: manufacturer’s instructions, safety warnings, and aftercare advice. This is part of designing responsibly, not an afterthought."
                ],
                "pitfalls": [
                    "Ignoring the user group (toy safety differs from power tool safety).",
                    "Vague warnings (must be specific and clear)."
                ],
                "exam_application": {
                    "prompt": "Suggest two safety warnings and one aftercare instruction for a consumer product.",
                    "what_to_do": [
                        "Make them specific and linked to hazards (heat, sharp edges, chemicals, choking).",
                        "Add aftercare to reduce risk over time (cleaning, inspection, replacement)."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.10",
        "title": "Protecting designs and intellectual property",
        "topics": [
            {
                "title": "IP protection: what each type protects and when you use it",
                "definitions": [
                    ("Copyright", "Protects original creative work (drawings, text, graphics)."),
                    ("Design rights", "Protects the appearance of a product (shape/configuration)."),
                    ("Patent", "Protects a new invention/technical solution (how it works)."),
                    ("Registered design", "Protects the visual appearance via formal registration."),
                    ("Trademark", "Protects brand identifiers (names, logos)."),
                    ("Logo", "A graphic brand mark (often protected by trademark and copyright)."),
                    ("Open design", "Sharing designs openly so others can use/adapt them (often with conditions).")
                ],
                "explanations": [
                    "Different protections cover different things: patents protect how it works, design rights/registered designs protect how it looks, and trademarks protect brand identity.",
                    "In exam answers, link the protection to the risk: copying the mechanism vs copying the appearance vs copying the brand.",
                    "Open design is the opposite approach: designs may be shared to encourage collaboration, innovation, or faster adoption - but you still need clear licensing."
                ],
                "pitfalls": [
                    "Saying ‘copyright protects inventions’ (that’s patents).",
                    "Mixing trademarks and registered designs (brand vs product appearance).",
                    "Ignoring that protection choices affect cost and business strategy."
                ],
                "exam_application": {
                    "prompt": "A company has created a new mechanism and a distinctive product shape. Which IP protections could they use?",
                    "what_to_do": [
                        "Patent for the mechanism; registered design/design rights for shape; trademark for name/logo.",
                        "Explain each briefly."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.11",
        "title": "Design for manufacturing, maintenance, repair and disposal",
        "topics": [
            {
                "title": "Designing for efficient manufacture (DFM)",
                "definitions": [
                    ("DFM", "Design for Manufacture: designing so it’s easier, faster and cheaper to make."),
                    ("Part count", "Number of separate parts; reducing it often reduces cost and errors."),
                    ("Standardised parts", "Common components/sizes used across products to reduce cost and simplify supply.")
                ],
                "explanations": [
                    "To make designs efficient, you reduce unnecessary processes and simplify assembly. Fewer steps usually means fewer errors and lower cost.",
                    "The spec also highlights using standardised parts and patterns/sizes to reduce waste and speed production.",
                    "For moulded products, design features like ribs/webbing add stiffness without lots of material. Snap-fits can remove the need for screws and speed assembly."
                ],
                "pitfalls": [
                    "Optimising for manufacturing but ruining user needs (must balance).",
                    "Adding ribs/snap-fits without considering moulding rules (draft, thickness, stress points).",
                    "Using custom parts where standard ones would work."
                ],
                "exam_application": {
                    "prompt": "Suggest three DFM changes to reduce the cost of a plastic product.",
                    "what_to_do": [
                        "Reduce part count; add ribs instead of thick walls; use snap-fits; standardise screws/components; add texture to avoid extra finishing."
                    ]
                }
            },
            {
                "title": "Maintenance and repair (designing so products can be serviced)",
                "definitions": [
                    ("Temporary fixing", "Can be undone for repair (screws, bolts)."),
                    ("Integral fixing", "Built-in joining features (snap-fits) that may allow disassembly."),
                    ("Serviceability", "How easy it is to access and replace parts.")
                ],
                "explanations": [
                    "The spec expects you to consider fixings, standardised parts, and allowing for service and repair/replacement of parts.",
                    "Designing for maintenance can extend product life. Sometimes products are upgraded through software downloads, which can improve performance without replacing hardware."
                ],
                "pitfalls": [
                    "Using permanent adhesives where repair access is needed.",
                    "Not thinking about tools/access panels/fastener types.",
                    "Forgetting that software upgrades still need user support and compatibility."
                ],
                "exam_application": {
                    "prompt": "Explain two design features that make a product easier to repair.",
                    "what_to_do": [
                        "Use screws/standard fasteners; modular sub-assemblies; labelled parts; accessible covers; replaceable battery if relevant."
                    ]
                }
            },
            {
                "title": "Disposal and sustainability (labelling, disassembly and the ‘six Rs’)",
                "definitions": [
                    ("Design for disassembly", "Designing so parts can be separated for repair/reuse/recycling."),
                    ("Labelling for recycling", "Marking materials so they can be sorted and recycled correctly."),
                    ("Reduce", "Use less material/energy and avoid harmful materials."),
                    ("Reuse", "Design parts to be used again (modules/components)."),
                    ("Rethink", "Choose eco-friendlier alternatives and smarter design approaches."),
                    ("Recycle", "Allow materials/components to be processed into new products.")
                ],
                "explanations": [
                    "The spec wants you to explain how material choice affects use, care and disposal. Key actions include labelling materials, making products easy to disassemble, and applying the sustainability ‘Rs’ listed in the specification (reduce, reuse, rethink, recycle).",
                    "Practical examples: avoid mixed materials that cannot be separated, avoid permanent laminates when recycling is a priority, and use fixings that can be undone so parts can be replaced instead of throwing away the whole product."
                ],
                "pitfalls": [
                    "Claiming recycling is easy when the design uses mixed materials and adhesives.",
                    "Not mentioning disassembly (recycling often depends on separation).",
                    "Forgetting that durability can be part of sustainability (longer life, less replacement)."
                ],
                "exam_application": {
                    "prompt": "Suggest two design changes that improve end-of-life for a product.",
                    "what_to_do": [
                        "Use separable fixings (screws/snap-fits), label materials, reduce coatings, standardise parts for reuse/replacement."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.12",
        "title": "Feasibility studies",
        "topics": [
            {
                "title": "Feasibility studies and prototype testing with users",
                "definitions": [
                    ("Feasibility study", "Checking if a design is practical to make, sell and use safely."),
                    ("Prototype", "A test version of a product used to learn and improve."),
                    ("User trial", "Testing with potential consumers to collect feedback and data.")
                ],
                "explanations": [
                    "Feasibility studies help you decide if an idea can realistically be produced. This includes checking materials, processes, cost, safety, and user acceptance.",
                    "Prototype testing with potential consumers is a key part: you collect feedback, measure performance, and refine the design before full production.",
                    "In exam questions, you normally explain what you would test (fit, comfort, strength, usability), how you would test it (tasks, measurements, questionnaires), and what decisions the results would influence."
                ],
                "pitfalls": [
                    "Only testing opinions (‘Do you like it?’) and not performance/usability.",
                    "Not testing with the actual target user group.",
                    "Ignoring cost/manufacturing feasibility (feasibility is not just ‘does it work’)."
                ],
                "exam_application": {
                    "prompt": "Describe a feasibility study for a new product concept.",
                    "what_to_do": [
                        "Include: prototype build, user testing, manufacturing checks, cost estimate, safety considerations.",
                        "Explain how results decide whether to progress or redesign."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.13",
        "title": "Enterprise and marketing in the development of products",
        "topics": [
            {
                "title": "Marketing and brand identity (how products are positioned and sold)",
                "definitions": [
                    ("Customer identification", "Defining your target market (who will buy/use it)."),
                    ("Brand identity", "The look/values/style that makes the brand recognisable."),
                    ("Labelling", "Information on the product/packaging (legal, safety, brand, usage)."),
                    ("Packaging", "Protects the product and communicates brand and information."),
                    ("Corporate identification", "Consistent brand visuals across products (logo, colours, typography).")
                ],
                "explanations": [
                    "Marketing is about making sure the right product reaches the right people in the right way. It starts with customer identification: who needs it, why, and what they value.",
                    "Labelling and packaging are not just decoration - they carry safety information, legal requirements, and brand messages. Corporate identification keeps everything consistent so the brand is instantly recognisable."
                ],
                "pitfalls": [
                    "Thinking marketing is only advertising (it includes target market, pricing, packaging and brand).",
                    "Ignoring legal/safety labelling requirements."
                ],
                "exam_application": {
                    "prompt": "Explain how packaging can support both function and marketing.",
                    "what_to_do": [
                        "Function: protection, storage, transport, instructions.",
                        "Marketing: brand identity, shelf impact, customer information, premium feel."
                    ]
                }
            },
            {
                "title": "Global marketing, promotion and new technologies",
                "definitions": [
                    ("Global marketing", "Promoting/selling across different countries/markets."),
                    ("Promotion", "Ways to get attention and persuade customers."),
                    ("Viral marketing", "Content shared rapidly by users through networks."),
                    ("Social media marketing", "Using platforms to reach and engage customers.")
                ],
                "explanations": [
                    "Global marketing means adapting messages to different cultures, languages and regulations while keeping brand identity consistent.",
                    "The spec specifically mentions promotion and advertisement using new technologies such as social media and viral marketing. In exam answers, you explain why these can be effective (targeting, low cost, sharing) and the risks (brand damage, misinformation, fast backlash)."
                ],
                "pitfalls": [
                    "Assuming the same campaign works everywhere (cultural differences matter).",
                    "Ignoring reputation risk (viral can go negative)."
                ],
                "exam_application": {
                    "prompt": "Give one advantage and one risk of using social media to promote a product.",
                    "what_to_do": [
                        "Advantage: targeted reach, low cost, quick feedback.",
                        "Risk: negative feedback spreads quickly; brand control is limited."
                    ]
                }
            },
            {
                "title": "Product costing, profit and entrepreneurs",
                "definitions": [
                    ("Costing", "Working out total cost of making and selling a product."),
                    ("Profit", "Money left after costs (profit = revenue - costs)."),
                    ("Entrepreneur", "Someone who starts and grows a business by taking risks.")
                ],
                "explanations": [
                    "Product costing includes materials, labour, overheads, tooling, waste, distribution and marketing. Profit only exists once all costs are covered.",
                    "Entrepreneurs spot opportunities, test ideas, manage risk, and drive innovation. In exam answers, link entrepreneurship to product development: new ideas, new markets, and problem solving."
                ],
                "pitfalls": [
                    "Only counting material cost (ignoring labour, overheads, tooling).",
                    "Confusing revenue and profit.",
                    "Not explaining what entrepreneurs actually do beyond ‘start a business’."
                ],
                "exam_application": {
                    "prompt": "A product sells for £30 and costs £18 to make. Calculate profit per unit and state one other cost a business must consider.",
                    "what_to_do": [
                        "Profit per unit: £12 (before other costs).",
                        "Mention overheads, marketing, distribution, packaging, tooling, returns."
                    ]
                }
            },
            {
                "title": "Collaborative working in product development",
                "definitions": [
                    ("Collaborative working", "Designers working together (in person or virtual) to develop products."),
                    ("Virtual collaboration", "Working together using shared files, video calls, cloud tools.")
                ],
                "explanations": [
                    "Modern product development is team-based. Collaboration speeds up problem solving and helps coordinate design, engineering and manufacturing.",
                    "Virtual and face-to-face systems are used to share CAD files, prototypes, test data and feedback. Good collaboration reduces mistakes and rework."
                ],
                "pitfalls": [
                    "Forgetting file versions (old CAD files cause manufacturing errors).",
                    "Vague answer: ‘teamwork is good’ (you must explain how it helps speed/quality/innovation)."
                ],
                "exam_application": {
                    "prompt": "Explain two benefits of collaborative working when developing a new product.",
                    "what_to_do": [
                        "Benefits: faster iteration, more expertise, fewer errors, better communication with manufacture and marketing."
                    ]
                }
            }
        ]
    },

    {
        "section": "3.1.14",
        "title": "Design communication",
        "topics": [
            {
                "title": "Communication and presentation techniques (what you must be able to do)",
                "definitions": [
                    ("Report writing", "Clear written communication: purpose, method, results, conclusions."),
                    ("Graph", "Visual display of data to show trends/relationships."),
                    ("Table / chart", "Organised data display for comparison."),
                    ("2D sketching", "Quick drawings to show ideas and details."),
                    ("3D sketching", "Drawings that show form/shape in 3D (depth)."),
                    ("Mixed media and rendering", "Using different materials/techniques to show realistic appearance."),
                    ("Dimensioning", "Adding sizes so the design can be made accurately."),
                    ("Details for manufacture", "Information like materials, tolerances, finishes, assembly notes."),
                    ("Scaling", "Drawing at a size ratio (eg 1:2, 2:1) while keeping proportions."),
                    ("Datum points", "Reference points/edges used to set out and measure accurately.")
                ],
                "explanations": [
                    "Design communication is about making sure your idea can be understood by clients, users and manufacturers - not just looking nice.",
                    "Reports explain decisions and evaluation. Graphs and charts show research results and test data clearly. Sketching shows ideas quickly; rendering shows materials/finishes and helps people ‘see’ the final product.",
                    "Dimensioning and manufacturing details turn ideas into buildable instructions. Scaling, geometry and datum points help accuracy when setting out drawings."
                ],
                "pitfalls": [
                    "Beautiful sketches with no dimensions or manufacturing detail.",
                    "Graphs without labels/units (no marks if the data is unclear).",
                    "Rendering that hides key features (communication must stay clear)."
                ],
                "exam_application": {
                    "prompt": "Explain two ways you would communicate a design to a manufacturer and why each is useful.",
                    "what_to_do": [
                        "Working drawings with dimensions/tolerances (accuracy).",
                        "Material/finish specification and assembly notes (correct manufacture).",
                        "Optional: graphs/tables from testing to justify choices."
                    ]
                }
            }
        ]
    }
]

# ---------- DOCX helpers ----------

def set_document_defaults(doc: Document, font_name="Calibri", font_size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    # ensure East Asian font mapping too
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), font_name)

def set_a4_margins(doc: Document, margin_cm=1.6):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)

def add_title_page(doc: Document, variant_name: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AQA A-Level Design & Technology: Product Design (7552)\n")
    run.bold = True
    run.font.size = Pt(20)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Paper 1 - Technical Principles (Section 3.1)\n")
    r2.bold = True
    r2.font.size = Pt(16)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Study Guide - {variant_name}\n")
    r3.font.size = Pt(14)
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("Based on the AQA specification (Version 1.4, 27 June 2022)\n")
    meta_run.italic = True
    meta_run.font.size = Pt(10)
    doc.add_paragraph()
    tips = doc.add_paragraph()
    tips_run = tips.add_run("How to use: Learn definitions, then practise applying them to exam-style scenarios.")
    tips_run.font.size = Pt(10)
    doc.add_page_break()

def add_section_header(doc: Document, sec_code: str, title: str):
    h = doc.add_heading(f"{sec_code}  {title}", level=1)
    return h

def add_topic_heading(doc: Document, title: str, level=2):
    return doc.add_heading(title, level=level)

def add_bullets(paragraph, items):
    # Not used directly; we'll add list paragraphs in doc
    pass

def add_two_col_card_table(doc: Document, defs, expls, pitfalls, exam_prompt, exam_steps, font_size=10.5):
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    labels = ["Core definitions (mark scheme)", "Explain it simply (but properly)", "Common pitfalls", "Exam application"]
    contents = []

    # Definitions text
    def_text = []
    for term, meaning in defs:
        def_text.append(f"{term}: {meaning}")
    contents.append(def_text)

    # Explanations (paragraphs)
    contents.append(expls)

    # Pitfalls
    contents.append(pitfalls)

    # Exam application
    ex = [f"Example question: {exam_prompt}", "How to answer (what to include):"] + [f"- {s}" for s in exam_steps]
    contents.append(ex)

    for r in range(4):
        # Label cell
        cell_label = table.cell(r, 0)
        cell_label.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell_label.paragraphs[0]
        p.clear()
        run = p.add_run(labels[r])
        run.bold = True
        run.font.size = Pt(font_size)

        # Content cell
        cell = table.cell(r, 1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        # Clear existing
        cell.paragraphs[0].clear()
        first = True
        for item in contents[r]:
            if item is None:
                continue
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            run = p.add_run(str(item))
            run.font.size = Pt(font_size)
            # convert "- " starter into bullet style for readability
            if isinstance(item, str) and item.strip().startswith("- "):
                p.style = "List Bullet"
                p.runs[0].text = item.strip()[2:]
        # If definitions/pitfalls, use bullets
        if r in (0,2):
            # replace paragraphs with bullets
            # remove all current paragraphs except first
            while len(cell.paragraphs) > 1:
                p = cell.paragraphs[-1]
                p._element.getparent().remove(p._element)
            cell.paragraphs[0].clear()
            for idx, item in enumerate(contents[r]):
                p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
                p.style = "List Bullet"
                run = p.add_run(str(item))
                run.font.size = Pt(font_size)
        if r == 1:
            # explanations as normal paragraphs
            pass
        if r == 3:
            # keep as mixed paragraphs; bullet step lines already handled
            pass

    # Column widths
    for row in table.rows:
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(15.0)

    doc.add_paragraph()  # spacing

def add_textbook_block(doc: Document, defs, expls, pitfalls, exam_prompt, exam_steps):
    # Definitions
    doc.add_paragraph("Core definitions (mark scheme)", style="Heading 3")
    for term, meaning in defs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{term}: ").bold = True
        p.add_run(meaning)

    # Explanations
    doc.add_paragraph("Detailed explanations (down-to-earth)", style="Heading 3")
    for para in expls:
        doc.add_paragraph(para)

    # Pitfalls
    doc.add_paragraph("Common pitfalls", style="Heading 3")
    for pit in pitfalls:
        doc.add_paragraph(pit, style="List Bullet")

    # Exam application
    doc.add_paragraph("Exam application", style="Heading 3")
    doc.add_paragraph(f"Example question: {exam_prompt}")
    doc.add_paragraph("What a strong answer includes:", style="Intense Quote")
    for step in exam_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph()

# ---------- Build documents ----------

out_dir = "/mnt/data"
docx_a = os.path.join(out_dir, "AQA_7552_Paper1_Technical_Principles_OptionA_Revision_Cards.docx")
docx_b = os.path.join(out_dir, "AQA_7552_Paper1_Technical_Principles_OptionB_Full_Guide.docx")

# Option A
docA = Document()
set_document_defaults(docA, font_size_pt=10.5)
set_a4_margins(docA, margin_cm=1.4)
add_title_page(docA, "Option A (Revision-card style)")

for s in guide:
    add_section_header(docA, s["section"], s["title"])
    # quick "what to know" line
    intro = docA.add_paragraph()
    intro.add_run("Focus: ").bold = True
    intro.add_run("Learn the terms, then practise linking them to real product needs, processes, and end-of-life.")
    for t in s["topics"]:
        add_topic_heading(docA, t["title"], level=2)
        ea = t["exam_application"]
        add_two_col_card_table(
            docA,
            t["definitions"],
            t["explanations"],
            t["pitfalls"],
            ea["prompt"],
            ea["what_to_do"],
            font_size=10
        )
    docA.add_page_break()

docA.save(docx_a)

# Option B
docB = Document()
set_document_defaults(docB, font_size_pt=11)
set_a4_margins(docB, margin_cm=1.6)
add_title_page(docB, "Option B (Full explanations)")

for s in guide:
    add_section_header(docB, s["section"], s["title"])
    docB.add_paragraph("This section is assessed through exam questions that test knowledge AND application. Use the examples to practise linking theory to product decisions.")
    for t in s["topics"]:
        add_topic_heading(docB, t["title"], level=2)
        ea = t["exam_application"]
        add_textbook_block(
            docB,
            t["definitions"],
            t["explanations"],
            t["pitfalls"],
            ea["prompt"],
            ea["what_to_do"]
        )
    docB.add_page_break()

docB.save(docx_b)

# ---------- Convert to PDFs ----------

def docx_to_pdf(docx_path: str):
    outdir = os.path.dirname(docx_path)
    cmd = [
        "soffice",
        f"-env:UserInstallation=file:///tmp/lo_profile_{os.getpid()}",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        outdir,
        docx_path
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)

docx_to_pdf(docx_a)
docx_to_pdf(docx_b)

pdf_a = docx_a.replace(".docx", ".pdf")
pdf_b = docx_b.replace(".docx", ".pdf")

(docx_a, pdf_a, docx_b, pdf_b)

